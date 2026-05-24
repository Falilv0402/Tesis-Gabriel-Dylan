"""
Genera RESUMEN_DOCUMENTACION_SARA.docx compilando los 5 archivos .md del proyecto.
Ejecutar: python docs/generar_resumen_word.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─── Colores del sistema SARA ────────────────────────────────────────────────
AZUL_OSCURO  = RGBColor(0x1e, 0x3a, 0x5f)   # títulos principales
AZUL_MEDIO   = RGBColor(0x2d, 0x6a, 0x9f)   # subtítulos
VERDE_SARA   = RGBColor(0x1a, 0x7a, 0x4a)   # etiquetas positivas
NARANJA_SARA = RGBColor(0xd9, 0x6b, 0x00)   # alertas / énfasis
GRIS_TABLA   = RGBColor(0xf0, 0xf4, 0xf8)   # fondo filas alternas
AZUL_HEADER  = RGBColor(0x1e, 0x3a, 0x5f)   # cabeceras tabla
BLANCO       = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_bg(cell, hex_color: str):
    """Aplica color de fondo a una celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(table):
    """Aplica bordes sutiles a toda la tabla."""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ["top", "left", "bottom", "right"]:
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "DADDE2")
                tcBorders.append(border)
            tcPr.append(tcBorders)


def add_heading(doc, text, level=1):
    """Heading estilizado con color."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(17)
        run.font.color.rgb = AZUL_OSCURO
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = AZUL_MEDIO
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = NARANJA_SARA
    return p


def add_body(doc, text, bold=False, italic=False, color=None):
    """Párrafo de cuerpo estándar."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(doc, text, level=0):
    """Bullet point."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def add_kpi_table(doc, rows_data, headers):
    """Tabla con cabecera azul oscuro y filas alternas."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Cabecera
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = BLANCO
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, "1e3a5f")

    # Filas de datos
    for idx, row_data in enumerate(rows_data):
        row = table.add_row()
        bg = "f0f4f8" if idx % 2 == 0 else "ffffff"
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_bg(cell, bg)

    set_cell_border(table)
    doc.add_paragraph()
    return table


def add_divider(doc):
    """Línea divisoria horizontal."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2d6a9f")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_badge(doc, text, color_hex):
    """Párrafo con texto destacado tipo badge."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"  {text}  ")
    run.bold = True
    run.font.size = Pt(9)
    r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
    run.font.color.rgb = RGBColor(r, g, b)
    return p


# ═══════════════════════════════════════════════════════════════════════════════
#  CONSTRUCCIÓN DEL DOCUMENTO
# ═══════════════════════════════════════════════════════════════════════════════
doc = Document()

# Márgenes
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Fuente base
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

# ─── PORTADA ─────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run("SARA")
run.font.size  = Pt(48)
run.bold       = True
run.font.color.rgb = AZUL_OSCURO

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Sistema de Alerta de Riesgo Académico")
r2.font.size  = Pt(18)
r2.font.color.rgb = AZUL_MEDIO

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Resumen General de Documentación Técnica")
r3.font.size  = Pt(13)
r3.italic     = True
r3.font.color.rgb = RGBColor(0x55, 0x65, 0x75)

doc.add_paragraph()
p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run(
    f"Proyecto P20261012  ·  UPC — Ingeniería de Sistemas de Información\n"
    f"Lima, Mayo 2026"
)
r4.font.size  = Pt(10.5)
r4.font.color.rgb = RGBColor(0x77, 0x88, 0x99)

doc.add_paragraph()
add_divider(doc)

# ─── INTRODUCCIÓN ────────────────────────────────────────────────────────────
doc.add_paragraph()
add_heading(doc, "¿Qué es este documento?", level=1)
add_body(doc,
    "Este documento compila en un solo lugar los 5 archivos de documentación técnica del proyecto SARA. "
    "Cada sección corresponde a un archivo .md del repositorio y presenta un resumen claro de su propósito, "
    "contenido principal y relevancia para la tesis. No es necesario leer los archivos originales para "
    "entender el alcance del proyecto; este Word sirve como punto de entrada unificado."
)
doc.add_paragraph()

add_kpi_table(doc,
    rows_data=[
        ["PLAN_SPRINTS_P20261012.md",         "Hoja de ruta",      "Sprints, HUs, arquitectura, KPIs del proyecto"],
        ["IMPLEMENTACION_HUS_CASOS.md",       "Cobertura técnica", "Mapeo épicas → capas del sistema + casos de prueba"],
        ["JUSTIFICACION_MODELO_PREDICTIVO.md","Fundamento ML",     "Por qué SARA es un modelo predictivo riguroso"],
        ["MODEL_CARD.md",                     "Ficha del modelo",  "Métricas, fairness, calibración, limitaciones"],
        ["ANALISIS_ERRORES.md",               "Análisis de fallos","Por qué y cómo se equivoca el modelo"],
    ],
    headers=["Archivo", "Tipo", "Qué contiene"]
)

add_divider(doc)
doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  DOCUMENTO 1 — PLAN DE SPRINTS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1 · PLAN_SPRINTS_P20261012.md", level=1)
add_body(doc, "Archivo: docs/PLAN_SPRINTS_P20261012.md", italic=True, color=RGBColor(0x88,0x88,0x88))
add_body(doc,
    "Es la hoja de ruta completa del proyecto. Define cómo se organiza el trabajo en 4 sprints, "
    "qué se construye en cada uno, quién es responsable, y cómo se mide el éxito al final. "
    "Es el documento que respondería la pregunta: '¿cómo planificaste tu tesis?'"
)
doc.add_paragraph()

add_heading(doc, "Contenido principal", level=2)

add_heading(doc, "Equipo y estructura", level=3)
add_bullet(doc, "Proyecto: P20261012 — Curso Taller de Proyectos I, UPC")
add_bullet(doc, "PM: Torres Saldaña, Gabriel Alonso · Scrum Manager: Tong Barahona, Dylan")
add_bullet(doc, "35 Historias de Usuario (HUs) · 73 escenarios de prueba · 7 Épicas")

add_heading(doc, "Sprints (4 en total)", level=3)
add_kpi_table(doc,
    rows_data=[
        ["Sprint 1", "Sem 4–6", "Modelo ML, dataset EM 2022, calibración, SHAP, métricas"],
        ["Sprint 2", "Sem 7–9", "Backend FastAPI (ML) + Backend principal (Supabase + Auth)"],
        ["Sprint 3", "Sem 10–12","Frontend Next.js: dashboard, predicciones, intervenciones"],
        ["Sprint 4", "Sem 13–15","Features avanzadas, despliegue cloud, sustentación final"],
    ],
    headers=["Sprint", "Semanas", "Foco"]
)

add_heading(doc, "Arquitectura del sistema", level=3)
add_kpi_table(doc,
    rows_data=[
        ["Frontend",       "Next.js 14 + TypeScript + Tailwind",      "Vistas del Director y Administrador"],
        ["Backend ML",     "FastAPI (Python)",                         "/predecir, /metricas, /diagnostico"],
        ["Backend Data",   "Supabase (Postgres + Auth + Storage)",     "7 tablas, RLS, roles, audit log"],
        ["Capa ML",        "LightGBM · XGBoost · SHAP · Optuna",      "Pipeline predictivo completo"],
        ["Cloud",          "Vercel · Railway · GitHub Actions · Sentry","CI/CD, monitoreo, despliegue"],
    ],
    headers=["Capa", "Tecnología", "Responsabilidad"]
)

add_heading(doc, "KPIs del proyecto (metas alcanzadas)", level=3)
add_kpi_table(doc,
    rows_data=[
        ["AUC-ROC test",     "≥ 0.80",  "0.843 ✅"],
        ["Recall (ALTO)",    "≥ 0.70",  "0.779 ✅"],
        ["Brier Score",      "≤ 0.18",  "0.157 ✅"],
        ["ECE calibración",  "≤ 0.08",  "0.048 ✅"],
        ["MCC",              "≥ 0.45",  "0.521 ✅"],
        ["HUs completadas",  "35/35",   "35/35 ✅"],
        ["Cobertura tests",  "≥ 80%",   "88% ✅"],
    ],
    headers=["Métrica / Entregable", "Meta", "Resultado"]
)

add_divider(doc)
doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  DOCUMENTO 2 — IMPLEMENTACIÓN HUs
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2 · IMPLEMENTACION_HUS_CASOS.md", level=1)
add_body(doc, "Archivo: docs/IMPLEMENTACION_HUS_CASOS.md", italic=True, color=RGBColor(0x88,0x88,0x88))
add_body(doc,
    "Es el puente entre el papel (las HUs del backlog) y el código real. Muestra qué épica "
    "del Product Backlog fue implementada en qué carpeta del repositorio, y qué casos de prueba "
    "validan cada funcionalidad. Es la respuesta técnica a '¿cómo se verifica que todo funciona?'"
)
doc.add_paragraph()

add_heading(doc, "Contenido principal", level=2)

add_heading(doc, "Cobertura por épica", level=3)
add_kpi_table(doc,
    rows_data=[
        ["EP01 — Acceso y Seguridad",         "Supabase Auth, roles, RLS, frontend login",          "CP001–CP010"],
        ["EP02 — Predicción y Clasificación",  "FastAPI /predecir, /predecir-dataset, niveles riesgo","CP011–CP020"],
        ["EP03 — Análisis y Visualización",    "Dashboard Next.js, KPIs, gráficos, heatmap",         "CP021–CP040"],
        ["EP04 — Priorización e Intervención", "Tabla intervenciones, relación estudiante/predicción","CP041–CP060"],
        ["EP05 — Histórico y Reportes",        "Predicciones versionadas, tablas por periodo",        "CP061–CP065"],
        ["EP06 — Gestión de Datos",            "Tablas académicas, Storage, dataset EM 2022",         "CP066–CP070"],
        ["EP07 — Mantenimiento ML",            "/metricas, /importancia, /reentrenar, artefactos",   "CP071–CP077"],
    ],
    headers=["Épica", "Implementación", "Casos de prueba"]
)

add_heading(doc, "Criterio de aceptación del primer incremento", level=3)
add_bullet(doc, "legacy-streamlit corre como demo funcional")
add_bullet(doc, "backend-ml responde health, métricas, importancia y predicciones")
add_bullet(doc, "supabase/migrations/0001_init.sql crea las 7 tablas base con RLS activo")
add_bullet(doc, "frontend compila y muestra la matriz de cumplimiento HU001–HU035 y CP001–CP077")

add_divider(doc)
doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  DOCUMENTO 3 — JUSTIFICACIÓN MODELO PREDICTIVO
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3 · JUSTIFICACION_MODELO_PREDICTIVO.md", level=1)
add_body(doc, "Archivo: docs/JUSTIFICACION_MODELO_PREDICTIVO.md", italic=True, color=RGBColor(0x88,0x88,0x88))
add_body(doc,
    "Es el documento más académico. Su propósito es demostrar, con rigor teórico y evidencia empírica, "
    "por qué SARA califica como un modelo predictivo de investigación. Incluye comparativa de algoritmos, "
    "descripción de todas las features, pipeline completo y respuestas preparadas para el jurado. "
    "Es la respuesta a '¿por qué esto es ciencia y no solo un script?'"
)
doc.add_paragraph()

add_heading(doc, "Contenido principal", level=2)

add_heading(doc, "¿Por qué SARA es un modelo predictivo?", level=3)
add_body(doc,
    "SARA satisface las 3 condiciones formales de la literatura (Bishop, 2006; Hastie et al., 2009):"
)
add_bullet(doc, "Generalización real: evaluado en IEs completamente nuevas (GroupShuffleSplit por ID_IE), AUC 0.843 en ~820 estudiantes de ~30 colegios nunca vistos")
add_bullet(doc, "Variable objetivo no disponible al momento de decidir: el modelo infiere el riesgo en Matemática antes de que el puntaje M500_M exista en el sistema")
add_bullet(doc, "Superación estadísticamente significativa de baselines (McNemar p<0.05, DeLong p<0.01): +63 pp en F1 respecto a la mejor heurística trivial")

add_heading(doc, "Comparativa de algoritmos (GroupKFold-5)", level=3)
add_kpi_table(doc,
    rows_data=[
        ["Logistic Regression", "0.865 ± 0.016", "Ganador por AUC-CV, modelo final"],
        ["LightGBM",            "0.855 ± 0.018", "Restricciones monotónicas, monotone_constraints"],
        ["Random Forest",       "0.855 ± 0.022", "n_estimators=300, monotonic_cst"],
        ["XGBoost",             "0.855 ± 0.017", "scale_pos_weight, monotone_constraints"],
        ["VotingClassifier",    "0.861 ± 0.014", "Soft-voting top-2, menor varianza"],
        ["Optuna (40 trials)",  "0.862 ± 0.015", "Bayesian tuning TPESampler"],
        ["FLAML AutoML (90s)",  "0.847 ± —",      "Cota superior de referencia"],
    ],
    headers=["Algoritmo", "AUC-CV (media ± std)", "Notas"]
)

add_heading(doc, "12 features del modelo con restricciones monotónicas", level=3)
add_kpi_table(doc,
    rows_data=[
        ["Sexo (F/M)",          "Categórica",  "Ninguna",  "Factor demográfico"],
        ["Nivel modalidad IE",  "Categórica",  "Ninguna",  "Pública / privada"],
        ["M500_L",              "Numérica",    "−1",       "Puntaje Lectura: mayor → menos riesgo"],
        ["M500_CN",             "Numérica",    "−1",       "Puntaje Ciencias: mayor → menos riesgo"],
        ["M500_M",              "Numérica",    "−1",       "Puntaje Matemática (variable objetivo proxy)"],
        ["ISE",                 "Numérica",    "−1",       "Índice socioeconómico: mayor → menos riesgo"],
        ["M500_L_iemean",       "Numérica IE", "−1",       "Media de Lectura en la IE"],
        ["M500_CN_iemean",      "Numérica IE", "−1",       "Media de Ciencias en la IE"],
        ["ISE_iemean",          "Numérica IE", "−1",       "ISE medio de la IE"],
        ["M500_L_relativa",     "Derivada",    "0",        "Diferencia vs. media Lectura de la IE"],
        ["M500_CN_relativa",    "Derivada",    "0",        "Diferencia vs. media Ciencias de la IE"],
        ["ISE_relativo",        "Derivada",    "0",        "Diferencia vs. ISE medio de la IE"],
    ],
    headers=["Feature", "Tipo", "Restricción", "Significado"]
)

add_heading(doc, "Pipeline técnico completo", level=3)
add_kpi_table(doc,
    rows_data=[
        ["Preprocesamiento",    "OrdinalEncoder + StandardScaler vía ColumnTransformer"],
        ["Validación cruzada",  "GroupKFold(5) + GroupShuffleSplit — respeta separación por IE"],
        ["Calibración",         "CalibratedClassifierCV(method='isotonic', cv=5)"],
        ["Interpretabilidad",   "SHAP (global + local + interacciones), PDP top-3, permutación (20 rep.)"],
        ["Pruebas estadísticas","McNemar, DeLong bootstrap, Hosmer-Lemeshow (χ²), Nested CV"],
        ["Estabilidad",         "10 semillas × GroupShuffleSplit → AUC media/std/min/max"],
        ["Equidad (Fairness)",  "Bootstrap IC95% por subgrupo (sexo, ISE tercil) + ECE por subgrupo"],
        ["AutoML comparación",  "FLAML 90s budget como cota superior de referencia"],
    ],
    headers=["Componente", "Descripción"]
)

add_divider(doc)
doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  DOCUMENTO 4 — MODEL CARD
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "4 · MODEL_CARD.md", level=1)
add_body(doc, "Archivo: docs/MODEL_CARD.md", italic=True, color=RGBColor(0x88,0x88,0x88))
add_body(doc,
    "Es el DNI o pasaporte del modelo. Sigue el estándar internacional Mitchell et al. (2019) "
    "— Model Cards for Model Reporting, ACM FAccT. Documenta todo lo que alguien necesita saber "
    "antes de usar o confiar en el modelo: qué hace, con qué datos fue entrenado, qué tan bueno es, "
    "si es justo con todos los grupos, y cuándo se debería reentrenar."
)
doc.add_paragraph()

add_heading(doc, "Contenido principal", level=2)

add_heading(doc, "Detalles técnicos del modelo", level=3)
add_kpi_table(doc,
    rows_data=[
        ["Nombre",            "SARA-EM-v1"],
        ["Tipo",              "Clasificador binario supervisado (riesgo académico Matemática)"],
        ["Algoritmo final",   "Logistic Regression calibrado (CalibratedClassifierCV isotonic)"],
        ["Features",          "12 variables (2 categóricas + 7 numéricas + 3 derivadas relativas)"],
        ["Datos",             "Evaluación Muestral EM 2022, Lima Metropolitana, gestión privada, 2.° grado"],
        ["Tamaño dataset",    "7,429 estudiantes — 5,921 entrenamiento / 1,508 prueba (~80/20)"],
        ["Distribución clase","39.4% ALTO riesgo / 60.6% sin riesgo (desbalance moderado)"],
        ["Artefactos",        "model/modelo_em.pkl · model/metricas_em.pkl"],
    ],
    headers=["Campo", "Valor"]
)

add_heading(doc, "Métricas de rendimiento (test set)", level=3)
add_kpi_table(doc,
    rows_data=[
        ["AUC-ROC",      "0.843",  "[0.811–0.874]", "Discriminación global"],
        ["PR-AUC",       "0.782",  "[0.741–0.822]", "Rendimiento en clase minoritaria"],
        ["Recall ALTO",  "0.779",  "[0.738–0.818]", "Detección de estudiantes en riesgo"],
        ["Precision",    "0.623",  "—",              "Exactitud de las alertas emitidas"],
        ["F1-Score",     "0.693",  "—",              "Balance Precision-Recall"],
        ["Accuracy",     "0.773",  "—",              "Exactitud global"],
        ["Brier Score",  "0.157",  "—",              "Error probabilístico (↓ mejor)"],
        ["Log-Loss",     "0.481",  "—",              "Entropía cruzada (↓ mejor)"],
        ["MCC",          "0.521",  "—",              "Correlación Matthews (robusto al desbalance)"],
        ["Specificity",  "0.769",  "—",              "Tasa de verdaderos negativos (TN rate)"],
        ["ECE",          "0.048",  "—",              "Error de calibración esperado (↓ mejor)"],
        ["MCE",          "0.094",  "—",              "Error de calibración máximo"],
        ["AUC imparcial","0.831",  "± 0.019",        "Nested CV — sin sesgo de selección"],
    ],
    headers=["Métrica", "Valor", "IC95% / Std", "Qué mide"]
)

add_heading(doc, "Auditoría de equidad (Fairness)", level=3)
add_body(doc,
    "Se evalúa que el modelo no discrimine entre grupos. Los intervalos de confianza bootstrap "
    "(500 iteraciones) se solapan entre grupos, indicando que no hay sesgo estadísticamente significativo:"
)
add_kpi_table(doc,
    rows_data=[
        ["Mujeres",          "0.847", "[0.806–0.887]", "0.781", "0.051"],
        ["Hombres",          "0.839", "[0.796–0.881]", "0.776", "0.049"],
        ["ISE Tercil 1 (bajo)","0.831","[0.788–0.874]","0.772", "0.053"],
        ["ISE Tercil 2 (medio)","0.844","[0.801–0.886]","0.779","0.047"],
        ["ISE Tercil 3 (alto)","0.852","[0.811–0.893]","0.784", "0.044"],
    ],
    headers=["Subgrupo", "AUC", "IC95% AUC", "Recall", "ECE"]
)

add_heading(doc, "Umbrales de monitoreo (cuándo reentrenar)", level=3)
add_kpi_table(doc,
    rows_data=[
        ["AUC-ROC",     "< 0.80",  "Degradación significativa de discriminación"],
        ["Recall ALTO", "< 0.72",  "Pérdida de sensibilidad a casos de riesgo"],
        ["ECE",         "> 0.08",  "Descalibración de probabilidades"],
        ["Brier Score", "> 0.20",  "Error probabilístico inaceptable"],
    ],
    headers=["Métrica", "Alerta si...", "Significado"]
)

add_divider(doc)
doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  DOCUMENTO 5 — ANÁLISIS DE ERRORES
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "5 · ANALISIS_ERRORES.md", level=1)
add_body(doc, "Archivo: docs/ANALISIS_ERRORES.md", italic=True, color=RGBColor(0x88,0x88,0x88))
add_body(doc,
    "Responde la pregunta más incómoda del comité: '¿cuándo y por qué falla tu modelo?' "
    "No basta con que las métricas sean buenas en promedio — hay que entender los casos "
    "específicos donde el modelo se equivoca, qué perfil tienen esos estudiantes y qué se puede "
    "hacer para mejorar. Es el documento que prueba que el análisis fue honesto y profundo."
)
doc.add_paragraph()

add_heading(doc, "Contenido principal", level=2)

add_heading(doc, "Tipos de error y su costo", level=3)
add_kpi_table(doc,
    rows_data=[
        ["Falso Negativo (FN)", "Estudiante en riesgo no detectado",
         "CRÍTICO: no recibe apoyo oportuno", "5×"],
        ["Falso Positivo (FP)", "Estudiante sin riesgo marcado incorrectamente",
         "Moderado: intervención innecesaria", "1×"],
    ],
    headers=["Tipo", "Qué significa", "Impacto", "Costo"]
)
add_body(doc,
    "Función de costo del pipeline: Costo_total = 5 × |FN| + 1 × |FP|. "
    "Esta asimetría justifica el umbral min_cost que privilegia Recall sobre Precision.",
    italic=True, color=RGBColor(0x55,0x66,0x77)
)

add_heading(doc, "Perfil del error más común", level=3)
add_bullet(doc, "FN típico: estudiante con M500_L entre 440–470 (zona gris), IE con ISE medio que 'nivela hacia abajo' las métricas relativas")
add_bullet(doc, "FP típico: estudiante con puntaje alto en Lectura/Ciencias pero ISE muy bajo (el modelo sobreestima el riesgo socioeconómico)")
add_bullet(doc, "El efecto de nivelación institucional (IEs con muchos estudiantes de riesgo similar) reduce la señal diferencial de las features relativas")

add_heading(doc, "Hallazgos del estudio de ablación (cuál feature reduce más los errores)", level=3)
add_kpi_table(doc,
    rows_data=[
        ["M500_L",            "−0.061", "La más importante — sin ella el AUC cae 6 puntos"],
        ["ISE",               "−0.038", "Segunda en importancia"],
        ["M500_CN",           "−0.029", "Tercera en importancia"],
        ["M500_L_relativa",   "−0.018", "Nueva feature derivada — impacto relevante"],
        ["M500_CN_relativa",  "−0.012", "Nueva feature derivada"],
        ["ISE_relativo",      "−0.008", "Nueva feature derivada — menor pero significativa"],
    ],
    headers=["Feature eliminada", "Delta AUC", "Interpretación"]
)
add_body(doc,
    "Las 3 variables relativas (rendimiento del estudiante vs. su propio colegio) redujeron "
    "los FN en aproximadamente 12% frente al modelo sin ellas.",
    italic=True, color=VERDE_SARA
)

add_heading(doc, "Distribución de errores por sexo (con bootstrap IC95%)", level=3)
add_kpi_table(doc,
    rows_data=[
        ["Mujeres",  "18.8%", "[15.2%–22.4%]", "21.0%", "[17.3%–24.7%]"],
        ["Hombres",  "19.4%", "[15.8%–23.0%]", "20.5%", "[16.9%–24.1%]"],
    ],
    headers=["Grupo", "Tasa FN", "IC95% FN", "Tasa FP", "IC95% FP"]
)
add_body(doc,
    "Los intervalos de confianza se solapan → no existe diferencia estadísticamente significativa "
    "en la tasa de error entre hombres y mujeres.",
    italic=True, color=VERDE_SARA
)

add_heading(doc, "Calibración por subgrupo (equidad probabilística)", level=3)
add_body(doc,
    "El ECE mide si las probabilidades del modelo son honestas para cada grupo. "
    "Un ECE bajo y similar entre grupos indica que el modelo no sobreestima ni subestima "
    "sistemáticamente el riesgo de ningún colectivo:"
)
add_kpi_table(doc,
    rows_data=[
        ["Mujeres",           "0.051", "0.098"],
        ["Hombres",           "0.049", "0.091"],
        ["ISE Tercil 1 bajo", "0.053", "0.104"],
        ["ISE Tercil 2 medio","0.047", "0.089"],
        ["ISE Tercil 3 alto", "0.044", "0.086"],
    ],
    headers=["Subgrupo", "ECE (calibración)", "MCE (peor bin)"]
)

add_heading(doc, "Checklist de mejoras", level=3)
add_bullet(doc, "✅ Feature engineering relativo (M500_L_relativa, M500_CN_relativa, ISE_relativo)")
add_bullet(doc, "✅ Monotonicity constraints en 6 features")
add_bullet(doc, "✅ Bootstrap IC95% global y por subgrupo")
add_bullet(doc, "✅ ECE por subgrupo (equidad de calibración)")
add_bullet(doc, "✅ Ablación de features (estudio completo 12 features)")
add_bullet(doc, "✅ Hosmer-Lemeshow test (χ² goodness-of-fit)")
add_bullet(doc, "✅ Nested CV (AUC imparcial sin sesgo de selección)")
add_bullet(doc, "✅ Stability analysis (10 semillas, AUC mean/std/min/max)")
add_bullet(doc, "⏳ Incorporar datos EM 2023 para validación longitudinal")
add_bullet(doc, "⏳ Variables conductuales (asistencia, participación) — no disponibles aún")
add_bullet(doc, "⏳ Extensión a colegios de gestión pública")
add_bullet(doc, "⏳ Modelo de supervivencia para proyección temporal real")

add_divider(doc)
doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  CIERRE
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Resumen final", level=1)
add_body(doc,
    "SARA es un sistema de alerta temprana de riesgo académico que cubre el ciclo completo "
    "de un proyecto de Machine Learning aplicado a educación: desde la planificación ágil en sprints "
    "hasta la auditoría ética del modelo, pasando por comparativa de algoritmos, calibración probabilística, "
    "interpretabilidad multicapa y pruebas estadísticas formales."
)
doc.add_paragraph()
add_kpi_table(doc,
    rows_data=[
        ["Pipeline ML",        "12 features · 7 algoritmos comparados · calibración isotónica · 20+ métricas"],
        ["Validación",         "GroupKFold-5 · GroupShuffleSplit · Nested CV · Hosmer-Lemeshow · McNemar · DeLong"],
        ["Interpretabilidad",  "SHAP global + local + interacciones · PDP · Permutation importance · Ablación"],
        ["Fairness",           "Bootstrap IC95% por subgrupo · ECE por subgrupo · sin sesgo sexo / ISE"],
        ["Sistema completo",   "FastAPI + Next.js 14 + Supabase + Leaflet heatmap Lima · 35 HUs · 73 CPs"],
    ],
    headers=["Dimensión", "Detalle"]
)
doc.add_paragraph()
add_body(doc,
    f"Documento generado automáticamente el {datetime.date.today().strftime('%d de %B de %Y')}. "
    "Para detalles completos, consultar los archivos .md individuales en docs/.",
    italic=True, color=RGBColor(0x88,0x88,0x88)
)

# ─── GUARDAR ─────────────────────────────────────────────────────────────────
output_path = r"C:\Users\Usuario\OneDrive\Escritorio\TesisDG-ML\PROYECTO-TESIS-DG\docs\RESUMEN_DOCUMENTACION_SARA.docx"
doc.save(output_path)
print(f"OK - Word generado en: {output_path}")
